'''
Created on 2019年10月6日

@author: liushucheng
'''
# coding=utf-8

import os
import re
import pysvn
import xlrd
import openpyxl
import string
import docx
from docx import Document

# 文档文书番号所在单元格的行列编号
row_num_doc_number = 6
col_num_doc_number = 4

# msn = ['adc','cortst','dio','eth','fls','flstst','gpt','icu','lin','mcu','port','pwm','ramtst','wdg']
# MSN = ['ADC','CORTST','DIO','ETH','FLS','FLSTST','GPT','ICU','LIN','MCU','PORT','PWM','RAMTST','WDG']

# MSN = ['DIO','FLS','MCU','PORT','WDG','ADC','CORTST','ETH','FLSTST','GPT','ICU','LIN','PWM','RAMTST']
# msn = ['dio','fls','mcu','port','wdg','adc','cortst','eth','flstst','gpt','icu','lin','pwm','ramtst']

# *********************************************************************************************************************
# Fetching UTP and UTR last changed version
# *********************************************************************************************************************
def FetchUTPUTRReport():
    print('****************************************************************************************************************')
    print('------------------------------------ Fetching UTP/UTR last changed revision ------------------------------------')
    print('****************************************************************************************************************')

    MSN = ['DIO','FLS','MCU','PORT','WDG','ADC','CORTST','ETH','FLSTST','GPT','ICU','LIN','PWM','RAMTST']
    msn = ['dio','fls','mcu','port','wdg','adc','cortst','eth','flstst','gpt','icu','lin','pwm','ramtst']

    client = pysvn.Client()
    for mod_index in range(len(MSN)):
        if MSN[mod_index] == 'ETH':
            path = "U:\\internal\\X1X\\F1x\\modules\\" + msn[mod_index] + "\\test_unit\\results\\Cantata\\"
        else:
            path = "U:\\internal\\X1X\\common_platform\\modules\\" + msn[mod_index] + "\\test_unit\\results\\Cantata\\"
        # 列出路径下的所有文件
        list_dir = os.listdir(path)
        for i in range(len(list_dir)):
            # print(list_dir[i])
            if 'UTR' in list_dir[i]:
                utr_srcfile = path + list_dir[i]
                # print(utr_srcfile)
                entry = client.info(utr_srcfile)
                try:
                    docx_file = Document(utr_srcfile)
                    # print(docx_file)
                    # print(str(len(docx_file.paragraphs)))
                    for para in docx_file.paragraphs:
                        # print(para.text)
                        if 'VERSION' in para.text:
                            # print(MSN[mod_index] + ',' + str(list_dir[i]) + ',' + para.text)
                            print(MSN[mod_index] + ',UTR,' + entry.url + ',VERSION: ' + para.text + ",svn: " + str(entry.commit_revision.number))
                except:
                    # print(MSN[mod_index] + ',' + str(list_dir[i]) + ',Note: *.DOC type can not be parsed')
                    print(MSN[mod_index] + ',UTR,' + entry.url + ",VERSION: *.DOC parse failed,svn: " + str(entry.commit_revision.number))
                # print(MSN[mod_index] + ',' + entry.name + ", UTR last changed revision, svn: " + str(entry.commit_revision.number))
                # print(MSN[mod_index] + ',UTR,' + entry.url + ",svn: " + str(entry.commit_revision.number))

    for mod_index in range(len(MSN)):
        if MSN[mod_index] == 'ETH':
            path = "U:\\internal\\X1X\\common_platform\\modules\\" + msn[mod_index] + "\\test_unit\\plan\\"
        else:
            path = "U:\\internal\\X1X\\common_platform\\modules\\" + msn[mod_index] + "\\test_unit\\plan\\Cantata\\"
        # 列出路径下的所有文件
        list_dir = os.listdir(path)
        for i in range(len(list_dir)):
            # print(list_dir[i])
            if 'UTP' in list_dir[i]:
                utp_srcfile = path + list_dir[i]
                # print(utr_srcfile)
                entry = client.info(utp_srcfile)
                try:
                    docx_file = Document(utp_srcfile)
                    # print(docx_file)
                    # print(str(len(docx_file.paragraphs)))
                    for para in docx_file.paragraphs:
                        # print(para.text)
                        if 'VERSION:' in para.text:
                            # print(MSN[mod_index] + ',' + str(list_dir[i]) + ',' + para.text)
                            print(MSN[mod_index] + ',UTP,' + entry.url + ',VERSION: ' + para.text + ",svn: " + str(entry.commit_revision.number))
                except:
                    # print(MSN[mod_index] + ',' + str(list_dir[i]) + ',Note: *.DOC type can not be parsed')
                    print(MSN[mod_index] + ',UTP,' + entry.url + ",VERSION: *.DOC parse failed,svn: " + str(entry.commit_revision.number))
                # print(MSN[mod_index] + ',' + entry.name + ", UTP last changed revision, svn: " + str(entry.commit_revision.number))
                # print(MSN[mod_index] + ',UTP,' + entry.url + ",svn: " + str(entry.commit_revision.number))

# *********************************************************************************************************************
# Fetching ESDD and TSDD last changed version
# *********************************************************************************************************************
def FetchESDDTSDDReport():
    print('****************************************************************************************************************')
    print('---------------------------------- Fetching ESDD/TSDD last changed revision ------------------------------------')
    print('****************************************************************************************************************')

    MSN = ['DIO','FLS','MCU','PORT','WDG','ADC','CORTST','ETH','FLSTST','GPT','ICU','LIN','PWM','RAMTST']
    msn = ['dio','fls','mcu','port','wdg','adc','cortst','eth','flstst','gpt','icu','lin','pwm','ramtst']

    client = pysvn.Client()

    for mod_index in range(len(MSN)):
        if MSN[mod_index] == 'ETH':
            path = "U:\\internal\\X1X\\common_platform\\modules\\" + msn[mod_index] + "\\docs\\sds\\"
        elif MSN[mod_index] == 'MCU':
            path = "U:\\internal\\X1X\\F1x\\modules\\" + msn[mod_index] + "\\docs\\sds\\"
        else:
            path = "U:\\internal\\X1X\\common_platform\\modules\\" + msn[mod_index] + "\\docs\\sds\\ESDD\\"
        # 列出路径下的所有文件
        list_dir = os.listdir(path)
        for i in range(len(list_dir)):
            # print(list_dir[i])
            if 'EAAR-SD-' in list_dir[i]:
                esdd_srcfile = path + list_dir[i]
                # print(esdd_srcfile)
                entry = client.info(esdd_srcfile)
                # print(MSN[mod_index] + ',' + entry.name + ",ESDD last changed revision: ,svn: " + str(entry.commit_revision.number))
                print(MSN[mod_index] + ',ESDD,' + entry.url + ",svn: " + str(entry.commit_revision.number))

    for mod_index in range(len(MSN)):
        if MSN[mod_index] == 'ETH':
            path = "U:\\internal\\X1X\\common_platform\\modules\\" + msn[mod_index] + "\\docs\\sds\\"
        elif MSN[mod_index] == 'MCU':
            path = "U:\\internal\\X1X\\F1x\\modules\\" + msn[mod_index] + "\\docs\\sds\\"
        else:
            path = "U:\\internal\\X1X\\common_platform\\modules\\" + msn[mod_index] + "\\docs\\sds\\TSDD\\"
        # 列出路径下的所有文件
        list_dir = os.listdir(path)
        for i in range(len(list_dir)):
            # print(list_dir[i])
            if 'EAAR-SD-' in list_dir[i]:
                tsdd_srcfile = path + list_dir[i]
                # print(esdd_srcfile)
                entry = client.info(tsdd_srcfile)
                # print(MSN[mod_index] + ',' + entry.name + ",TSDD last changed revision: ,svn: " + str(entry.commit_revision.number))
                print(MSN[mod_index] + ',TSDD,' + entry.url + ",svn: " + str(entry.commit_revision.number))

# *********************************************************************************************************************
# Fetching EUM and TUM last changed version
# *********************************************************************************************************************
def FetchUMReport():
    print('****************************************************************************************************************')
    print('------------------------------------ Fetching EUM/TUM last changed revision ------------------------------------')
    print('****************************************************************************************************************')

    MSN = ['DIO','FLS','MCU','PORT','WDG','ADC','CORTST','ETH','FLSTST','GPT','ICU','LIN','PWM','RAMTST']
    msn = ['dio','fls','mcu','port','wdg','adc','cortst','eth','flstst','gpt','icu','lin','pwm','ramtst']

    client = pysvn.Client()

    for mod_index in range(len(MSN)):
        # get last changed revision of CUM in external folder
        cum_srcfile_external = "U:\\external\\X1X\\F1x\\modules\\" + msn[mod_index] + "\\user_manual\\AUTOSAR_" + MSN[mod_index] + "_Component_UserManual.pdf"
        entry = client.info(cum_srcfile_external)
        # print(MSN[mod_index] + ',' + entry.name + ",EUM last changed revision, svn: " + str(entry.commit_revision.number))
        print(MSN[mod_index] + ',EUM,' + entry.url + ",svn: " + str(entry.commit_revision.number))

    for mod_index in range(len(MSN)):
        # get last changed revision of TUM in external folder
        tum_srcfile_external = "U:\\external\\X1X\\F1x\\modules\\" + msn[mod_index] + "\\user_manual\\AUTOSAR_" + MSN[mod_index] + "_Tool_UserManual.pdf"
        entry = client.info(tum_srcfile_external)
        # print(MSN[mod_index] + ',' + entry.name + ",TUM last changed revision, svn: " + str(entry.commit_revision.number))
        print(MSN[mod_index] + ',TUM,' + entry.url + ",svn: " + str(entry.commit_revision.number))

# *********************************************************************************************************************
# Fetching TEQ report last changed version
# *********************************************************************************************************************
def FetchTEQReport():
    print('****************************************************************************************************************')
    print('------------------------------------ Fetching TEQ last changed revision ----------------------------------------')
    print('****************************************************************************************************************')

    MSN = ['DIO','FLS','MCU','PORT','WDG','ADC','CORTST','ETH','FLSTST','GPT','ICU','LIN','PWM','RAMTST']
    msn = ['dio','fls','mcu','port','wdg','adc','cortst','eth','flstst','gpt','icu','lin','pwm','ramtst']

    client = pysvn.Client()

    TEQ_PATH = "U:\\internal\\X1X\\F1x\\common_family\\docs\\FuSa\\Tool_Evaluation_and_Qualification\\Code_Generator_Ver4.05.00.B_Ver42.05.00.B\\"
    # list all files
    list_teq = os.listdir(TEQ_PATH)
    # walk through files
    for i in range(len(list_teq)):
        # construct full path of files
        TEQ_SrcFile = TEQ_PATH + list_teq[i]
        # pysvn get last changed revision
        entry = client.info(TEQ_SrcFile)
        for j in range(len(MSN)):
            if MSN[j] in list_teq[i]:
                # print(str(entry.url) + ',' + entry.name + ',VERSION: undefined,svn: ' + str(entry.commit_revision.number))
                print(MSN[j] + ',TEQ,' + entry.url + ',VERSION: undefined,svn: ' + str(entry.commit_revision.number))
            else:
                pass

# *********************************************************************************************************************
# FMEA版本号以及last changed version信息读取
# *********************************************************************************************************************
def FetchFMEAReport():
    print('****************************************************************************************************************')
    print('----------------------------------- Fetching FMEA last changed revision ----------------------------------------')
    print('****************************************************************************************************************')

    MSN = ['DIO','FLS','MCU','PORT','WDG','ADC','CORTST','ETH','FLSTST','GPT','ICU','LIN','PWM','RAMTST']
    msn = ['dio','fls','mcu','port','wdg','adc','cortst','eth','flstst','gpt','icu','lin','pwm','ramtst']

    fmea_row_num_version = 12
    fmea_col_num_version = 4

    client = pysvn.Client()

    for mod_index in range(len(MSN)):
        if MSN[mod_index] == 'ETH':
            fmea_srcfile = "U:\\internal\\X1X\\F1x\\modules\\" + msn[mod_index] + "\\docs\\safety_analysis\\F1KM_Ver4.05.00_Ver42.05.00_F1KH_Ver42.05.00_ASILB\\F1KM_Ver4.05.00_Ver42.05.00_F1KH_Ver42.05.00_SafetyAnalysis_" + MSN[mod_index] + ".xlsx"
        else:
            fmea_srcfile = "U:\\internal\\X1X\\F1x\\modules\\" + msn[mod_index] + "\\docs\\safety_analysis\\F1K_F1KM_Ver4.05.00_Ver42.05.00_F1KH_Ver42.05.00_ASILB\\F1K_F1KM_Ver4.05.00_Ver42.05.00_F1KH_Ver42.05.00_SafetyAnalysis_" + MSN[mod_index] + ".xlsx"
        wb_fmea = openpyxl.load_workbook(fmea_srcfile,data_only = True)
        # print(wb_fmea)
        ws_names = wb_fmea.sheetnames
        # print(ws_names)
        ws_cover = wb_fmea[ws_names[0]]
        # print(ws_cover)
        doc_number = ws_cover.cell(row = row_num_doc_number,column = col_num_doc_number).value
        # print(MSN[mod_index] + ',' + fmea_srcfile + ', FMEA document number,' + doc_number)
        version = ws_cover.cell(row = fmea_row_num_version,column = fmea_col_num_version).value
        # print(MSN[mod_index] + ',' + fmea_srcfile + ', FMEA version,' + version)

        entry = client.info(fmea_srcfile)

        # print(MSN[mod_index] + ',' + entry.name + ',' + doc_number + ',' + version)
        print(MSN[mod_index] + ',FMEA,' + entry.url + ',' + doc_number + ',' + version)

# *********************************************************************************************************************
# DFA-CC 版本号以及last changed version信息读取
# *********************************************************************************************************************
def FetchDFAReport():
    print('****************************************************************************************************************')
    print('--------------------------------- Fetching DFA-CC last changed revision ----------------------------------------')
    print('****************************************************************************************************************')

    MSN = ['DIO','FLS','MCU','PORT','WDG','ADC','CORTST','ETH','FLSTST','GPT','ICU','LIN','PWM','RAMTST']
    msn = ['dio','fls','mcu','port','wdg','adc','cortst','eth','flstst','gpt','icu','lin','pwm','ramtst']
    # DFA
    dfa_row_num_version = 14
    dfa_col_num_version = 4

    client = pysvn.Client()

    for mod_index in range(len(MSN)):
        if MSN[mod_index] == 'ETH':
            dfa_srcfile = "U:\\internal\\X1X\\F1x\\modules\\" + msn[mod_index] + "\\docs\\safety_analysis\\F1KM_Ver4.05.00_Ver42.05.00_F1KH_Ver42.05.00_ASILB\\F1KM_Ver4.05.00_Ver42.05.00_F1KH_Ver42.05.00_DFA_CC_" + MSN[mod_index] + ".xlsm"
        else:
            dfa_srcfile = "U:\\internal\\X1X\\F1x\\modules\\" + msn[mod_index] + "\\docs\\safety_analysis\\F1K_F1KM_Ver4.05.00_Ver42.05.00_F1KH_Ver42.05.00_ASILB\\F1K_F1KM_Ver4.05.00_Ver42.05.00_F1KH_Ver42.05.00_DFA_CC_" + MSN[mod_index] + ".xlsm"
        wb_dfa = openpyxl.load_workbook(dfa_srcfile,data_only = True)
        # print(wb_fmea)
        ws_names = wb_dfa.sheetnames
        # print(ws_names)
        ws_cover = wb_dfa[ws_names[0]]
        # print(ws_cover)
        doc_number = ws_cover.cell(row = row_num_doc_number,column = col_num_doc_number).value
        # print(MSN[mod_index] + ',' + dfa_srcfile + ', DFA-CC document number,' + doc_number)
        version = ws_cover.cell(row = dfa_row_num_version,column = dfa_col_num_version).value
        # print(MSN[mod_index] + ',' + dfa_srcfile + ', DFA-CC version,' + version)

        entry = client.info(dfa_srcfile)

        # print(MSN[mod_index] + ',' + entry.name + ',' + doc_number + ',' + version)
        print(MSN[mod_index] + ',DFA,' + entry.url + ',' + doc_number + ',' + version)

# *********************************************************************************************************************
# ESTS 版本号以及last changed version信息读取
# *********************************************************************************************************************
def FetchESTSReport():
    print('****************************************************************************************************************')
    print('--------------------------------- Fetching ESTS last changed revision ------------------------------------------')
    print('****************************************************************************************************************')

    MSN = ['DIO','FLS','MCU','PORT','WDG','ADC','CORTST','ETH','FLSTST','GPT','ICU','LIN','PWM','RAMTST']
    msn = ['dio','fls','mcu','port','wdg','adc','cortst','eth','flstst','gpt','icu','lin','pwm','ramtst']

    # ESTS
    ests_row_num_version = 16
    ests_col_num_version = 1

    client = pysvn.Client()

    for mod_index in range(len(MSN)):
        try:
            if MSN[mod_index] == 'MCU':
                ests_srcfile = "U:\\internal\\X1X\\F1x\\modules\\" + msn[mod_index] + "\\test_func\\plan\\AUTOSAR_Renesas_" + MSN[mod_index] + "_ESTS_X1x.xlsx"
            else:
                ests_srcfile = "U:\\internal\\X1X\\common_platform\\modules\\" + msn[mod_index] + "\\test_func\\plan\\AUTOSAR_Renesas_" + MSN[mod_index] + "_ESTS_X1x.xlsx"
            wb_ests = openpyxl.load_workbook(ests_srcfile,data_only = True)
            ws_names = wb_ests.sheetnames
            ws_cover = wb_ests[ws_names[0]]
            version = ws_cover.cell(row = ests_row_num_version,column = ests_col_num_version).value
            # print(MSN[mod_index] + ', ESTS version,' + version)
            entry = client.info(ests_srcfile)
            # print(MSN[mod_index] + ',' + entry.name + ',' + version + ",svn: " + str(entry.commit_revision.number))
            print(MSN[mod_index] + ',ESTS,' + entry.url + ',' + version + ",svn: " + str(entry.commit_revision.number))
        except:
            print(MSN[mod_index] + ', ESTS does not exsit.')


# *********************************************************************************************************************
# ESTR 版本号以及last changed version信息读取
# *********************************************************************************************************************
def FetchESTRReport():
    print('****************************************************************************************************************')
    print('--------------------------------- Fetching ESTR last changed revision ------------------------------------------')
    print('****************************************************************************************************************')

    MSN = ['DIO','FLS','MCU','PORT','SPI','WDG','ADC','CORTST','ETH','FLSTST','GPT','ICU','LIN','PWM','RAMTST']
    msn = ['dio','fls','mcu','port','spi','wdg','adc','cortst','eth','flstst','gpt','icu','lin','pwm','ramtst']

    # ESTR openpyxl行列从1开始计算
    estr_row_num_version = 18
    estr_col_num_version = 1
    # ESTR xlrd行列从0开始计算
    estr_xls_row_num_version = 17
    estr_xls_col_num_version = 0

    client = pysvn.Client()

    for mod_index in range(len(MSN)):
        path_r403 = "U:\\internal\\X1X\\F1x\\modules\\" + msn[mod_index] + "\\test_func\\results\\4.0.3\\"
        path_r422 = "U:\\internal\\X1X\\F1x\\modules\\" + msn[mod_index] + "\\test_func\\results\\4.2.2\\"
        list_dir_r403 = os.listdir(path_r403)
        list_dir_r422 = os.listdir(path_r422)
        # 403 ESTR
        for i in range(len(list_dir_r403)):
            if '.xls' in list_dir_r403[i]:
                estr_srcfile_r403 = path_r403 + list_dir_r403[i]
                # print(estr_srcfile_r403)
                # xlsx文件处理
                try:
                    wb_estr_r403 = openpyxl.load_workbook(estr_srcfile_r403,data_only = True,read_only = True)
                    # print(wb_estr_r403)
                    ws_names_r403 = wb_estr_r403.sheetnames
                    # print(ws_names_r403)
                    ws_cover_r403 = wb_estr_r403[ws_names_r403[0]]
                    # print(ws_cover_r403)
                    version = ws_cover_r403.cell(row = estr_row_num_version,column = estr_col_num_version).value
                    # print(MSN[mod_index] + ',' + str(list_dir_r403[i]) + 'R403 ESTR version,' + version)
                except:
                    # print(MSN[mod_index],list_dir_r403[i],'R403 ACCESS FAILED')
                    xls_wbfile = xlrd.open_workbook(estr_srcfile_r403)
                    # print(xls_wbfile.nsheets)
                    # 通过sheet页的索引取得第一页的内容
                    xls_sheet = xls_wbfile.sheet_by_index(0)
                    # print(sheet.name,', total rows number =',sheet.nrows,', total cols number =',sheet.ncols)
                    # 从版本号的记录位置读取版本号信息
                    version = xls_sheet.cell(estr_xls_row_num_version,estr_xls_col_num_version).value
                    # print(MSN[mod_index] + ',' + str(list_dir_r403[i]) + 'R403 ESTR version,' + version)
                # pysvn get last changed revision
                entry = client.info(estr_srcfile_r403)
                print(MSN[mod_index] + ',ESTR,' + entry.name + ',' +  version + ',svn: ' + str(entry.commit_revision.number))
                # print(MSN[mod_index] + ',ESTR,' + entry.url + ',' +  version + ',svn: ' + str(entry.commit_revision.number))
        # 422 ESTR
        for i in range(len(list_dir_r422)):
            if '.xls' in list_dir_r422[i]:
                estr_srcfile_r422 = path_r422 + list_dir_r422[i]
                try:
                    wb_estr_r422 = openpyxl.load_workbook(estr_srcfile_r422,data_only = True,read_only = True)
                    ws_names_r422 = wb_estr_r422.sheetnames
                    ws_cover_r422 = wb_estr_r422[ws_names_r422[0]]
                    version = ws_cover_r422.cell(row = estr_row_num_version,column = estr_col_num_version).value
                    # print(MSN[mod_index] + ',' + str(list_dir_r422[i]) + 'R422 ESTR version,' + version)
                except:
                    # print(MSN[mod_index],list_dir_r422[i],'R422 ACCESS FAILED')
                    # 加载工作簿
                    xls_wbfile = xlrd.open_workbook(estr_srcfile_r422)
                    # print(xls_wbfile.nsheets)
                    # 通过sheet页的索引取得第一页的内容
                    xls_sheet = xls_wbfile.sheet_by_index(0)
                    # print(sheet.name,', total rows number =',sheet.nrows,', total cols number =',sheet.ncols)
                    # 从版本号的记录位置读取版本号信息
                    version = xls_sheet.cell(estr_xls_row_num_version,estr_xls_col_num_version).value
                    # print(MSN[mod_index] + ',' + str(list_dir_r422[i]) + 'R422 ESTR version,' + version)
                # pysvn get last changed revision
                entry = client.info(estr_srcfile_r422)
                print(MSN[mod_index] + ',ESTR,' + entry.name + ',' +  version + ',svn: ' + str(entry.commit_revision.number))
                # print(MSN[mod_index] + ',ESTR,' + entry.url + ',' +  version + ',svn: ' + str(entry.commit_revision.number))

# *********************************************************************************************************************
# TSTP版本号以及last changed version信息读取
# *********************************************************************************************************************
def FetchTSTPReport():
    print('****************************************************************************************************************')
    print('---------------------------- Fetching TSTP Report last changed revision ----------------------------------------')
    print('****************************************************************************************************************')

    MSN = ['DIO','FLS','MCU','PORT','WDG','ADC','CORTST','ETH','FLSTST','GPT','ICU','LIN','PWM','RAMTST']
    msn = ['dio','fls','mcu','port','wdg','adc','cortst','eth','flstst','gpt','icu','lin','pwm','ramtst']
    # TSTP openpyxl行列从1开始计算
    tstp_row_num_version = 14
    tstp_col_num_version = 1
    # TSTP xlrd行列从0开始计算
    tstp_xls_row_num_version = 13
    tstp_xls_col_num_version = 0

    client = pysvn.Client()

    for mod_index in range(len(MSN)):
        # 构造文件路径
        if MSN[mod_index] == 'MCU':
            tstp_path = "U:\\internal\\X1X\\F1x\\modules\\" + msn[mod_index] + "\\test_cfg\\plan\\"
        else:
            tstp_path = "U:\\internal\\X1X\\common_platform\\modules\\" + msn[mod_index] + "\\test_cfg\\plan\\"
        # 返回tstp_path指定的文件夹包含的文件或文件夹的名字的列表
        list_tstp = os.listdir(tstp_path)
        # 从返回的名字列表中提取名称里包含'.xls'的文件
        for i in range(len(list_tstp)):
            if '.xls' in list_tstp[i]:
                tstp_srcfile = tstp_path + list_tstp[i]
                # print(tstp_srcfile)
                # 扩展名是xlsx的文件尝试用openpyxl加载
                try:
                    wb_tstp = openpyxl.load_workbook(tstp_srcfile,data_only = True,read_only = True)
                    ws_tstp = wb_tstp.sheetnames
                    ws_tstp_cover = wb_tstp[ws_tstp[0]]
                    version = ws_tstp_cover.cell(row = tstp_row_num_version,column = tstp_col_num_version).value
                    # print(MSN[mod_index] + ',' + str(list_tstp[i]) + 'TSTP version,' + version)
                # 扩展名是xls的文件尝试用xlrd加载
                except:
                    wb_tstp = xlrd.open_workbook(tstp_srcfile)
                    ws_tstp_cover = wb_tstp.sheet_by_index(0)
                    # 从版本号的记录位置读取版本号信息
                    version = ws_tstp_cover.cell(tstp_xls_row_num_version,tstp_xls_col_num_version).value
                    # print(MSN[mod_index] + ',' + str(list_tstp[i]) + 'TSTP version,' + version)
                # pysvn get last changed revision
                entry = client.info(tstp_srcfile)
                # print(MSN[mod_index] + ',' + entry.name + "TSTP last changed revision," + str(entry.commit_revision.number))
                print(MSN[mod_index] + ',TSTP,' + entry.url + ',' + version + ',svn: ' + str(entry.commit_revision.number))

# *********************************************************************************************************************
# TSTR版本号以及last changed version信息读取
# *********************************************************************************************************************
def FetchTSTRReport():
    print('****************************************************************************************************************')
    print('------------------------------- Fetching TSTR last changed revision --------------------------------------------')
    print('****************************************************************************************************************')

    MSN = ['DIO','FLS','MCU','PORT','WDG','ADC','CORTST','ETH','FLSTST','GPT','ICU','LIN','PWM','RAMTST']
    msn = ['dio','fls','mcu','port','wdg','adc','cortst','eth','flstst','gpt','icu','lin','pwm','ramtst']
    # TSTR openpyxl行列从1开始计算
    tstr_row_num_version = 14
    tstr_col_num_version = 1
    # TSTP xlrd行列从0开始计算
    tstr_xls_row_num_version = 13
    tstr_xls_col_num_version = 0

    client = pysvn.Client()

    # TSTR R403
    for mod_index in range(len(MSN)):
        # 构造文件路径
        if MSN[mod_index] == 'MCU':
            tstr_path_r403 = "U:\\internal\\X1X\\F1x\\modules\\" + msn[mod_index] + "\\test_cfg\\results\\Autosar4.0.3\\"
        elif MSN[mod_index] in ['CORTST','FLSTST','PWM','RAMTST','FLS']:
            tstr_path_r403 = "U:\\internal\\X1X\\common_platform\\modules\\" + msn[mod_index] + "\\test_cfg\\results\\Autosar4.0.3\\"
        # ETH的目录结构不同于其他模块，ETH没有创建4.0.3和4.2.2目录，只有一个results目录，该目录下的两个Excel是F1KM 4.0.3和4.2.2的报告
        elif MSN[mod_index] == 'ETH':
            tstr_path_r403 = "U:\\internal\\X1X\\common_platform\\modules\\" + msn[mod_index] + "\\test_cfg\\results\\"
        else:
            tstr_path_r403 = "U:\\internal\\X1X\\common_platform\\modules\\" + msn[mod_index] + "\\test_cfg\\results\\4.0.3\\"
        # 返回tstr_path_r403指定的文件夹包含的文件或文件夹的名字的列表
        list_tstr = os.listdir(tstr_path_r403)
        for i in range(len(list_tstr)):
            if '.xls' in list_tstr[i]:
                tstr_r403_srcfile = tstr_path_r403 + list_tstr[i]
                # print(tstr_r403_srcfile)
                # 扩展名是xlsx的文件尝试用openpyxl加载
                try:
                    wb_tstr_r403 = openpyxl.load_workbook(tstr_r403_srcfile,data_only = True,read_only = True)
                    ws_tstr_r403 = wb_tstr_r403.sheetnames
                    ws_tstr_cover = wb_tstr_r403[ws_tstr_r403[0]]
                    version = ws_tstr_cover.cell(row = tstr_row_num_version,column = tstr_col_num_version).value
                    # print(MSN[mod_index] + ',' + str(list_tstr[i]) + "TSTR R403 version," + version)
                except:
                    wb_tstr_r403 = xlrd.open_workbook(tstr_r403_srcfile)
                    # ws_names = wb_tstr_r403.sheet_names()
                    ws_tstr_cover = wb_tstr_r403.sheet_by_index(0)
                    # 从版本号的记录位置读取版本号信息
                    version = ws_tstr_cover.cell(tstr_xls_row_num_version,tstr_xls_col_num_version).value
                    # print(MSN[mod_index] + ',' + str(list_tstr[i]) + "TSTR R403 version," + version)
                # pysvn get last changed revision
                entry = client.info(tstr_r403_srcfile)
                # print(MSN[mod_index] + ',' + entry.name + ',' + version + ',svn: ' + str(entry.commit_revision.number))
                print(MSN[mod_index] + ',TSTR,' + entry.url + ',' + version + ',svn: ' + str(entry.commit_revision.number))

    # TSTR R422
    for mod_index in range(len(MSN)):
        # 构造文件路径
        if MSN[mod_index] == 'MCU':
            tstr_path_r422 = "U:\\internal\\X1X\\F1x\\modules\\" + msn[mod_index] + "\\test_cfg\\results\\Autosar4.2.2\\"
        elif MSN[mod_index] in ['CORTST','FLSTST','PWM','RAMTST','FLS']:
            tstr_path_r422 = "U:\\internal\\X1X\\common_platform\\modules\\" + msn[mod_index] + "\\test_cfg\\results\\Autosar4.2.2\\"
        # ETH的目录结构不同于其他模块，ETH没有创建4.0.3和4.2.2目录，只有一个results目录，该目录下还有一个F1KH目录
        elif MSN[mod_index] == 'ETH':
            tstr_path_r422 = "U:\\internal\\X1X\\common_platform\\modules\\" + msn[mod_index] + "\\test_cfg\\results\\F1KH\\"
        else:
            tstr_path_r422 = "U:\\internal\\X1X\\common_platform\\modules\\" + msn[mod_index] + "\\test_cfg\\results\\4.2.2\\"
        # 返回tstr_path_r403指定的文件夹包含的文件或文件夹的名字的列表
        list_tstr = os.listdir(tstr_path_r422)
        for i in range(len(list_tstr)):
            if '.xls' in list_tstr[i]:
                tstr_r422_srcfile = tstr_path_r422 + list_tstr[i]
                # print(tstr_r403_srcfile)
                # 扩展名是xlsx的文件尝试用openpyxl加载
                try:
                    wb_tstr_r422 = openpyxl.load_workbook(tstr_r422_srcfile,data_only = True,read_only = True)
                    ws_tstr_r422 = wb_tstr_r422.sheetnames
                    ws_tstr_cover = wb_tstr_r422[ws_tstr_r403[0]]
                    version = ws_tstr_cover.cell(row = tstr_row_num_version,column = tstr_col_num_version).value
                    # print(MSN[mod_index] + ',' + str(list_tstr[i]) + "TSTR R422 version," + version)
                except:
                    wb_tstr_r422 = xlrd.open_workbook(tstr_r422_srcfile)
                    # ws_names = wb_tstr_r422.sheet_names()
                    ws_tstr_cover = wb_tstr_r422.sheet_by_index(0)
                    # 从版本号的记录位置读取版本号信息
                    version = ws_tstr_cover.cell(tstr_xls_row_num_version,tstr_xls_col_num_version).value
                    # print(MSN[mod_index] + ',' + str(list_tstr[i]) + "TSTR R422 version," + version)
                # pysvn get last changed revision
                entry = client.info(tstr_r422_srcfile)
                # print(MSN[mod_index] + ',' + entry.name + "TSTR R422 last changed revision, svn: " + str(entry.commit_revision.number))
                print(MSN[mod_index] + ',TSTR,' + entry.url + ',' + version + ',svn: ' + str(entry.commit_revision.number))

# *********************************************************************************************************************
# QAC 版本号以及last changed version信息读取
# *********************************************************************************************************************
def FetchQACReport():
    print('****************************************************************************************************************')
    print('--------------------------- Fetching QAC Report last changed revision ------------------------------------------')
    print('****************************************************************************************************************')

    MSN = ['DIO','FLS','MCU','PORT','WDG','ADC','CORTST','ETH','FLSTST','GPT','ICU','LIN','PWM','RAMTST']
    msn = ['dio','fls','mcu','port','wdg','adc','cortst','eth','flstst','gpt','icu','lin','pwm','ramtst']
    # QAC openpyxl行列从1开始计算
    qac_row_num_version = 18
    qac_col_num_version = 2
    # QAC xlrd行列从0开始计算
    qac_xls_row_num_version = 17
    qac_xls_col_num_version = 1

    client = pysvn.Client()

    # QAC R403
    for mod_index in range(len(MSN)):
        qac_path_r403 = "U:\\internal\\X1X\\F1x\\modules\\" + msn[mod_index] + "\\test_static\\qac\\4.0.3\\"
        list_qac = os.listdir(qac_path_r403)
        for i in range(len(list_qac)):
            if 'Metrics.xls' in list_qac[i]:
                qac_srcfile = qac_path_r403 + list_qac[i]
                try:
                    wb_qac = openpyxl.load_workbook(qac_srcfile,data_only = True,read_only = True)
                    ws_qac = wb_qac.sheetnames
                    ws_qac_cover = wb_qac[ws_qac[0]]
                    version = ws_qac_cover.cell(row = qac_row_num_version,column = qac_col_num_version).value
                    # print(MSN[mod_index] + ',' + str(list_qac[i]) + "QAC R403 version," + version)
                except:
                    wb_qac = xlrd.open_workbook(qac_srcfile)
                    ws_qac_cover = wb_qac.sheet_by_index(0)
                    # 从版本号的记录位置读取版本号信息
                    version = ws_qac_cover.cell(qac_xls_row_num_version,qac_xls_col_num_version).value
                    # print(MSN[mod_index] + ',' + str(list_qac[i]) + "QAC R403 version," + version)
                # pysvn get last changed revision
                entry = client.info(qac_srcfile)
                print(MSN[mod_index] + ',QAC,' + entry.name + ',' + version + ",svn: " + str(entry.commit_revision.number))
                # print(MSN[mod_index] + ',QAC,' + entry.url + ',' + version + ",svn: " + str(entry.commit_revision.number))
    # QAC R422
    for mod_index in range(len(MSN)):
        qac_path_r422 = "U:\\internal\\X1X\\F1x\\modules\\" + msn[mod_index] + "\\test_static\\qac\\4.2.2\\"
        list_qac = os.listdir(qac_path_r422)
        for i in range(len(list_qac)):
            if 'Metrics.xls' in list_qac[i]:
                qac_srcfile = qac_path_r422 + list_qac[i]
                try:
                    wb_qac = openpyxl.load_workbook(qac_srcfile,data_only = True,read_only = True)
                    ws_qac = wb_qac.sheetnames
                    ws_qac_cover = wb_qac[ws_qac[0]]
                    version = ws_qac_cover.cell(row = qac_row_num_version,column = qac_col_num_version).value
                    # print(MSN[mod_index] + ',' + str(list_qac[i]) + "QAC R422 version," + version)
                except:
                    wb_qac = xlrd.open_workbook(qac_srcfile)
                    ws_qac_cover = wb_qac.sheet_by_index(0)
                    # 从版本号的记录位置读取版本号信息
                    version = ws_qac_cover.cell(qac_xls_row_num_version,qac_xls_col_num_version).value
                    # print(MSN[mod_index] + ',' + str(list_qac[i]) + "QAC R422 version," + version)
                # pysvn get last changed revision
                entry = client.info(qac_srcfile)
                print(MSN[mod_index] + ',QAC,' + entry.name +',' + version + ",svn: " + str(entry.commit_revision.number))
                # print(MSN[mod_index] + ',QAC,' + entry.url + ',' + version + ",svn: " + str(entry.commit_revision.number))

# *********************************************************************************************************************
# Fetching Reqtify *.rqtf last changed version (Reqtify 报告中不包含version信息)
# *********************************************************************************************************************
def FetchReqtifyReport():
    print('****************************************************************************************************************')
    print('------------------------- Fetching Reqtify Report last changed revision ----------------------------------------')
    print('****************************************************************************************************************')

    MSN = ['DIO','FLS','MCU','PORT','WDG','ADC','CORTST','ETH','FLSTST','GPT','ICU','LIN','PWM','RAMTST']
    msn = ['dio','fls','mcu','port','wdg','adc','cortst','eth','flstst','gpt','icu','lin','pwm','ramtst']

    client = pysvn.Client()

    for mod_index in range(len(MSN)):
        if MSN[mod_index] in ['ADC','DIO','ETH','GPT','LIN','PORT','RAMTST']:
            reqtify_path = "U:\\internal\\X1X\\F1x\\modules\\" +  msn[mod_index] + "\\reqtify\\"
        else:
            reqtify_path = "U:\\internal\\X1X\\F1x\\modules\\" +  msn[mod_index] + "\\Reqtify\\"
        # 返回reqtify_path指定的文件夹包含的文件或文件夹的名字的列表
        list_reqtify_path = os.listdir(reqtify_path)
        for i in range(len(list_reqtify_path)):
            if list_reqtify_path[i] in ['F1K_403','F1K_422','F1KH_422','F1KM_403','F1KM_422']:
                reqtify_report_path = reqtify_path + list_reqtify_path[i]
                list_reqtify_report_path = os.listdir(reqtify_report_path)
                for j in range(len(list_reqtify_report_path)):
                    # regex 匹配字符结尾
                    res = re.search('.+rqtf$',list_reqtify_report_path[j])
                    if res != None:
                        # print(res)
                        # print(list_reqtify_report_path[j])
                        reqtify_srcfile = reqtify_report_path + "\\" + list_reqtify_report_path[j]
                        # print(reqtify_srcfile)
                        # pysvn get last changed revision
                        entry = client.info(reqtify_srcfile)
                        # print(MSN[mod_index] + ',' + str(list_reqtify_path[i]) + ',' + entry.name + ',svn: ' + str(entry.commit_revision.number))
                        print(MSN[mod_index] + ',Reqtify,' + entry.url + ',VERSION: undefined,svn: ' + str(entry.commit_revision.number))

# *********************************************************************************************************************
# Fetching Reqtify Uncovered_Requirements_Justification report last changed version (Reqtify报告中不包含version信息)
# *********************************************************************************************************************
def FetchReqtifyUncoveredRequirementsJustificationReport():
    print('****************************************************************************************************************')
    print('---------------- Fetching Uncovered_Requirements_Justification Report last changed revision --------------------')
    print('****************************************************************************************************************')

    # ETH 模块的目录结构不方便程序化，重新定义了模块列表，把ETH模块从列表中移除
    MSN = ['DIO','FLS','MCU','PORT','WDG','ADC','CORTST','FLSTST','GPT','ICU','LIN','PWM','RAMTST']
    msn = ['dio','fls','mcu','port','wdg','adc','cortst','flstst','gpt','icu','lin','pwm','ramtst']

    client = pysvn.Client()

    for mod_index in range(len(MSN)):
        if MSN[mod_index] in ['ADC','DIO','ETH','GPT','LIN','PORT','RAMTST']:
            reqtify_path = "U:\\internal\\X1X\\F1x\\modules\\" +  msn[mod_index] + "\\reqtify\\"
        else:
            reqtify_path = "U:\\internal\\X1X\\F1x\\modules\\" +  msn[mod_index] + "\\Reqtify\\"
        # 返回reqtify_path指定的文件夹包含的文件或文件夹的名字的列表
        list_reqtify_path = os.listdir(reqtify_path)
        for i in range(len(list_reqtify_path)):
            if list_reqtify_path[i] in ['F1K_403']:
                if MSN[mod_index] in ['ADC','DIO','ICU','GPT','PWM']:
                    reqtify_report_path = reqtify_path + list_reqtify_path[i] + "\\Reports\\F1K\\Uncovered_Requirements\\"
                    list_reqtify_report_path = os.listdir(reqtify_report_path)
                    for j in range(len(list_reqtify_report_path)):
                        # 由于DIO模块的Justification文件名中的Justificaiton拼写错误，所以为了通配，关键字只用了部分字母
                        if 'fication' in list_reqtify_report_path[j]:
                            reqtify_srcfile = reqtify_report_path + list_reqtify_report_path[j]
                            # pysvn get last changed revision
                            entry = client.info(reqtify_srcfile)
                            # print(MSN[mod_index] + ',' + str(list_reqtify_path[i]) + ',' + entry.name + ',svn: ' + str(entry.commit_revision.number))
                            print(MSN[mod_index] + ',ReqtifyJustification,' + entry.url + ',svn: ' + str(entry.commit_revision.number))
                else:
                    reqtify_report_path = reqtify_path + list_reqtify_path[i] + "\\Reports\\Uncovered_Requirements\\"
                    list_reqtify_report_path = os.listdir(reqtify_report_path)
                    for j in range(len(list_reqtify_report_path)):
                        # 由于DIO模块的Justification文件名中的Justificaiton拼写错误，所以为了通配，关键字只用了部分字母
                        if 'fication' in list_reqtify_report_path[j]:
                            reqtify_srcfile = reqtify_report_path + list_reqtify_report_path[j]
                            # pysvn get last changed revision
                            entry = client.info(reqtify_srcfile)
                            # print(MSN[mod_index] + ',' + str(list_reqtify_path[i]) + ',' + entry.name + ',svn: ' + str(entry.commit_revision.number))
                            print(MSN[mod_index] + ',ReqtifyJustification,' + entry.url + ',svn: ' + str(entry.commit_revision.number))
            elif list_reqtify_path[i] in ['F1K_422']:
                if MSN[mod_index] in ['ADC','DIO','ICU','PWM']:
                    reqtify_report_path = reqtify_path + list_reqtify_path[i] + "\\Reports\\F1K\\Uncovered_Requirements\\"
                    list_reqtify_report_path = os.listdir(reqtify_report_path)
                    for j in range(len(list_reqtify_report_path)):
                        # 由于DIO模块的Justification文件名中的Justificaiton拼写错误，所以为了通配，关键字只用了部分字母
                        if 'fication' in list_reqtify_report_path[j]:
                            reqtify_srcfile = reqtify_report_path + list_reqtify_report_path[j]
                            # pysvn get last changed revision
                            entry = client.info(reqtify_srcfile)
                            # print(MSN[mod_index] + ',' + str(list_reqtify_path[i]) + ',' + entry.name + ',svn: ' + str(entry.commit_revision.number))
                            print(MSN[mod_index] + ',ReqtifyJustification,' + entry.url + ',svn: ' + str(entry.commit_revision.number))
                else:
                    reqtify_report_path = reqtify_path + list_reqtify_path[i] + "\\Reports\\Uncovered_Requirements\\"
                    list_reqtify_report_path = os.listdir(reqtify_report_path)
                    for j in range(len(list_reqtify_report_path)):
                        # 由于DIO模块的Justification文件名中的Justificaiton拼写错误，所以为了通配，关键字只用了部分字母
                        if 'fication' in list_reqtify_report_path[j]:
                            reqtify_srcfile = reqtify_report_path + list_reqtify_report_path[j]
                            # pysvn get last changed revision
                            entry = client.info(reqtify_srcfile)
                            # print(MSN[mod_index] + ',' + str(list_reqtify_path[i]) + ',' + entry.name + ',svn: ' + str(entry.commit_revision.number))
                            print(MSN[mod_index] + ',ReqtifyJustification,' + entry.url + ',svn: ' + str(entry.commit_revision.number))
            elif list_reqtify_path[i] in ['F1KM_403','F1KM_422']:
                if MSN[mod_index] in ['ADC','DIO','ICU','GPT','PWM']:
                    reqtify_report_path = reqtify_path + list_reqtify_path[i] + "\\Reports\\F1KM\\Uncovered_Requirements\\"
                    list_reqtify_report_path = os.listdir(reqtify_report_path)
                    for j in range(len(list_reqtify_report_path)):
                        # 由于DIO模块的Justification文件名中的Justificaiton拼写错误，所以为了通配，关键字只用了部分字母
                        if 'fication' in list_reqtify_report_path[j]:
                            reqtify_srcfile = reqtify_report_path + list_reqtify_report_path[j]
                            # pysvn get last changed revision
                            entry = client.info(reqtify_srcfile)
                            # print(MSN[mod_index] + ',' + str(list_reqtify_path[i]) + ',' + entry.name + ',svn: ' + str(entry.commit_revision.number))
                            print(MSN[mod_index] + ',ReqtifyJustification,' + entry.url + ',svn: ' + str(entry.commit_revision.number))
                else:
                    reqtify_report_path = reqtify_path + list_reqtify_path[i] + "\\Reports\\Uncovered_Requirements\\"
                    list_reqtify_report_path = os.listdir(reqtify_report_path)
                    for j in range(len(list_reqtify_report_path)):
                        # 由于DIO模块的Justification文件名中的Justificaiton拼写错误，所以为了通配，关键字只用了部分字母
                        if 'fication' in list_reqtify_report_path[j]:
                            reqtify_srcfile = reqtify_report_path + list_reqtify_report_path[j]
                            # pysvn get last changed revision
                            entry = client.info(reqtify_srcfile)
                            # print(MSN[mod_index] + ',' + str(list_reqtify_path[i]) + ',' + entry.name + ',svn: ' + str(entry.commit_revision.number))
                            print(MSN[mod_index] + ',ReqtifyJustification,' + entry.url + ',svn: ' + str(entry.commit_revision.number))
            elif list_reqtify_path[i] in ['F1KH_422']:
                if MSN[mod_index] in ['ADC','DIO','ICU','GPT','PWM']:
                    reqtify_report_path = reqtify_path + list_reqtify_path[i] + "\\Reports\\F1KH\\Uncovered_Requirements\\"
                    list_reqtify_report_path = os.listdir(reqtify_report_path)
                    for j in range(len(list_reqtify_report_path)):
                        # 由于DIO模块的Justification文件名中的Justificaiton拼写错误，所以为了通配，关键字只用了部分字母
                        if 'fication' in list_reqtify_report_path[j]:
                            reqtify_srcfile = reqtify_report_path + list_reqtify_report_path[j]
                            # pysvn get last changed revision
                            entry = client.info(reqtify_srcfile)
                            # print(MSN[mod_index] + ',' + str(list_reqtify_path[i]) + ',' + entry.name + ',svn: ' + str(entry.commit_revision.number))
                            print(MSN[mod_index] + ',ReqtifyJustification,' + entry.url + ',svn: ' + str(entry.commit_revision.number))
                else:
                    reqtify_report_path = reqtify_path + list_reqtify_path[i] + "\\Reports\\Uncovered_Requirements\\"
                    list_reqtify_report_path = os.listdir(reqtify_report_path)
                    for j in range(len(list_reqtify_report_path)):
                        # 由于DIO模块的Justification文件名中的Justificaiton拼写错误，所以为了通配，关键字只用了部分字母
                        if 'fication' in list_reqtify_report_path[j]:
                            reqtify_srcfile = reqtify_report_path + list_reqtify_report_path[j]
                            # pysvn get last changed revision
                            entry = client.info(reqtify_srcfile)
                            # print(MSN[mod_index] + ',' + str(list_reqtify_path[i]) + ',' + entry.name + ',svn: ' + str(entry.commit_revision.number))
                            print(MSN[mod_index] + ',ReqtifyJustification,' + entry.url + ',svn: ' + str(entry.commit_revision.number))

# *********************************************************************************************************************
# Reqtify Traceability_Reports report last changed version信息读取 (Reqtify报告中不包含version信息)
# *********************************************************************************************************************
def FetchReqtifyTraceabilityReports():
    print('****************************************************************************************************************')
    print('------------------------ Fetching Traceability Reports last changed revision -----------------------------------')
    print('****************************************************************************************************************')

    # ETH 模块的目录结构不方便程序化，重新定义了模块列表，把ETH模块从列表中移除
    MSN = ['DIO','FLS','MCU','PORT','WDG','ADC','CORTST','FLSTST','GPT','ICU','LIN','PWM','RAMTST']
    msn = ['dio','fls','mcu','port','wdg','adc','cortst','flstst','gpt','icu','lin','pwm','ramtst']

    client = pysvn.Client()

    for mod_index in range(len(MSN)):
        if MSN[mod_index] in ['ADC','DIO','ETH','GPT','LIN','PORT','RAMTST']:
            reqtify_path = "U:\\internal\\X1X\\F1x\\modules\\" +  msn[mod_index] + "\\reqtify\\"
        else:
            reqtify_path = "U:\\internal\\X1X\\F1x\\modules\\" +  msn[mod_index] + "\\Reqtify\\"
        # 返回reqtify_path指定的文件夹包含的文件或文件夹的名字的列表
        list_reqtify_path = os.listdir(reqtify_path)
        for i in range(len(list_reqtify_path)):
            if list_reqtify_path[i] in ['F1K_403']:
                if MSN[mod_index] in ['ADC','DIO','ICU','GPT','PWM']:
                    reqtify_report_path = reqtify_path + list_reqtify_path[i] + "\\Reports\\F1K\\Traceability_Reports\\"
                    list_reqtify_report_path = os.listdir(reqtify_report_path)
                    for j in range(len(list_reqtify_report_path)):
                        # 由于DIO模块的Justification文件名中的Justificaiton拼写错误，所以为了通配，关键字只用了部分字母
                        # 由于ADC模块下的文件是without filter的类型,没有取到
                        if 'With_Filter' in list_reqtify_report_path[j]:
                            reqtify_srcfile = reqtify_report_path + list_reqtify_report_path[j]
                            # pysvn get last changed revision
                            entry = client.info(reqtify_srcfile)
                            # print(MSN[mod_index] + ',' + str(list_reqtify_path[i]) + ',' + entry.name + ',svn: ' + str(entry.commit_revision.number))
                            print(MSN[mod_index] + ',ReqtifyTraceability,' + entry.url + ',svn: ' + str(entry.commit_revision.number))
                else:
                    reqtify_report_path = reqtify_path + list_reqtify_path[i] + "\\Reports\\Traceability_Reports\\"
                    list_reqtify_report_path = os.listdir(reqtify_report_path)
                    for j in range(len(list_reqtify_report_path)):
                        # 由于DIO模块的Justification文件名中的Justificaiton拼写错误，所以为了通配，关键字只用了部分字母
                        if 'With_Filter' in list_reqtify_report_path[j]:
                            reqtify_srcfile = reqtify_report_path + list_reqtify_report_path[j]
                            # pysvn get last changed revision
                            entry = client.info(reqtify_srcfile)
                            # print(MSN[mod_index] + ',' + str(list_reqtify_path[i]) + ',' + entry.name + ',svn: ' + str(entry.commit_revision.number))
                            print(MSN[mod_index] + ',ReqtifyTraceability,' + entry.url + ',svn: ' + str(entry.commit_revision.number))
            elif list_reqtify_path[i] in ['F1K_422']:
                if MSN[mod_index] in ['ADC','DIO','ICU','PWM']:
                    reqtify_report_path = reqtify_path + list_reqtify_path[i] + "\\Reports\\F1K\\Traceability_Reports\\"
                    list_reqtify_report_path = os.listdir(reqtify_report_path)
                    for j in range(len(list_reqtify_report_path)):
                        # 由于DIO模块的Justification文件名中的Justificaiton拼写错误，所以为了通配，关键字只用了部分字母
                        if 'With_Filter' in list_reqtify_report_path[j]:
                            reqtify_srcfile = reqtify_report_path + list_reqtify_report_path[j]
                            # pysvn get last changed revision
                            entry = client.info(reqtify_srcfile)
                            # print(MSN[mod_index] + ',' + str(list_reqtify_path[i]) + ',' + entry.name + ',svn: ' + str(entry.commit_revision.number))
                            print(MSN[mod_index] + ',ReqtifyTraceability,' + entry.url + ',svn: ' + str(entry.commit_revision.number))
                else:
                    reqtify_report_path = reqtify_path + list_reqtify_path[i] + "\\Reports\\Traceability_Reports\\"
                    list_reqtify_report_path = os.listdir(reqtify_report_path)
                    for j in range(len(list_reqtify_report_path)):
                        # 由于DIO模块的Justification文件名中的Justificaiton拼写错误，所以为了通配，关键字只用了部分字母
                        if 'With_Filter' in list_reqtify_report_path[j]:
                            reqtify_srcfile = reqtify_report_path + list_reqtify_report_path[j]
                            # pysvn get last changed revision
                            entry = client.info(reqtify_srcfile)
                            # print(MSN[mod_index] + ',' + str(list_reqtify_path[i]) + ',' + entry.name + ',svn: ' + str(entry.commit_revision.number))
                            print(MSN[mod_index] + ',ReqtifyTraceability,' + entry.url + ',svn: ' + str(entry.commit_revision.number))
            elif list_reqtify_path[i] in ['F1KM_403','F1KM_422']:
                if MSN[mod_index] in ['ADC','DIO','ICU','GPT','PWM']:
                    reqtify_report_path = reqtify_path + list_reqtify_path[i] + "\\Reports\\F1KM\\Traceability_Reports\\"
                    list_reqtify_report_path = os.listdir(reqtify_report_path)
                    for j in range(len(list_reqtify_report_path)):
                        # 由于DIO模块的Justification文件名中的Justificaiton拼写错误，所以为了通配，关键字只用了部分字母
                        if 'With_Filter' in list_reqtify_report_path[j]:
                            reqtify_srcfile = reqtify_report_path + list_reqtify_report_path[j]
                            # pysvn get last changed revision
                            entry = client.info(reqtify_srcfile)
                            # print(MSN[mod_index] + ',' + str(list_reqtify_path[i]) + ',' + entry.name + ',svn: ' + str(entry.commit_revision.number))
                            print(MSN[mod_index] + ',ReqtifyTraceability,' + entry.url + ',svn: ' + str(entry.commit_revision.number))
                else:
                    reqtify_report_path = reqtify_path + list_reqtify_path[i] + "\\Reports\\Traceability_Reports\\"
                    list_reqtify_report_path = os.listdir(reqtify_report_path)
                    for j in range(len(list_reqtify_report_path)):
                        # 由于DIO模块的Justification文件名中的Justificaiton拼写错误，所以为了通配，关键字只用了部分字母
                        if 'With_Filter' in list_reqtify_report_path[j]:
                            reqtify_srcfile = reqtify_report_path + list_reqtify_report_path[j]
                            # pysvn get last changed revision
                            entry = client.info(reqtify_srcfile)
                            # print(MSN[mod_index] + ',' + str(list_reqtify_path[i]) + ',' + entry.name + ',svn: ' + str(entry.commit_revision.number))
                            print(MSN[mod_index] + ',ReqtifyTraceability,' + entry.url + ',svn: ' + str(entry.commit_revision.number))
            elif list_reqtify_path[i] in ['F1KH_422']:
                if MSN[mod_index] in ['ADC','DIO','ICU','GPT','PWM']:
                    reqtify_report_path = reqtify_path + list_reqtify_path[i] + "\\Reports\\F1KH\\Traceability_Reports\\"
                    list_reqtify_report_path = os.listdir(reqtify_report_path)
                    for j in range(len(list_reqtify_report_path)):
                        # 由于DIO模块的Justification文件名中的Justificaiton拼写错误，所以为了通配，关键字只用了部分字母
                        if 'With_Filter' in list_reqtify_report_path[j]:
                            reqtify_srcfile = reqtify_report_path + list_reqtify_report_path[j]
                            # pysvn get last changed revision
                            entry = client.info(reqtify_srcfile)
                            # print(MSN[mod_index] + ',' + str(list_reqtify_path[i]) + ',' + entry.name + ',svn: ' + str(entry.commit_revision.number))
                            print(MSN[mod_index] + ',ReqtifyTraceability,' + entry.url + ',svn: ' + str(entry.commit_revision.number))
                else:
                    reqtify_report_path = reqtify_path + list_reqtify_path[i] + "\\Reports\\Traceability_Reports\\"
                    list_reqtify_report_path = os.listdir(reqtify_report_path)
                    for j in range(len(list_reqtify_report_path)):
                        # 由于DIO模块的Justification文件名中的Justificaiton拼写错误，所以为了通配，关键字只用了部分字母
                        if 'With_Filter' in list_reqtify_report_path[j]:
                            reqtify_srcfile = reqtify_report_path + list_reqtify_report_path[j]
                            # pysvn get last changed revision
                            entry = client.info(reqtify_srcfile)
                            # print(MSN[mod_index] + ',' + str(list_reqtify_path[i]) + ',' + entry.name + ',svn: ' + str(entry.commit_revision.number))
                            print(MSN[mod_index] + ',ReqtifyTraceability,' + entry.url + ',svn: ' + str(entry.commit_revision.number))

# *********************************************************************************************************************
# Fetching Peer Review Minutes report last changed version (confirm only the doc number at cover sheet)
# *********************************************************************************************************************
def FetchPeerReviewMinutes():
    print('****************************************************************************************************************')
    print('------------------ Fetching Peer Review Minutes report last changed revision -----------------------------------')
    print('****************************************************************************************************************')

    MSN = ['DIO','FLS','MCU','PORT','WDG','ADC','CORTST','ETH','FLSTST','GPT','ICU','LIN','PWM','RAMTST','GENERAL']
    msn = ['dio','fls','mcu','port','wdg','adc','cortst','eth','flstst','gpt','icu','lin','pwm','ramtst','general']

    # document number location U3
    row_num_document_number = 3
    col_num_document_number = 21

    client = pysvn.Client()

    for mod_index in range(len(MSN)):
        peer_review_minutes_path = "U:\\internal\\X1X\\F1x\\modules\\" + msn[mod_index] + "\\review\\ILCD\\F1K_F1KM_Ver4.05.00_Ver42.05.00_F1KH_Ver42.05.00_ASILB\\"
        list_peer_review_minutes = os.listdir(peer_review_minutes_path)
        for i in range(len(list_peer_review_minutes)):
            # 获得review议事录路径
            peer_reivew_minutes_srcfile = peer_review_minutes_path + list_peer_review_minutes[i]
            # 加载工作簿
            wb_peer_review_minutes_srcfile = openpyxl.load_workbook(peer_reivew_minutes_srcfile,data_only = True,read_only = True)
            # 取得工作簿中所有工作表
            ws_peer_review_minutes_srcfile = wb_peer_review_minutes_srcfile.sheetnames
            # 根据工作表索引取得封面页
            ws_cover = wb_peer_review_minutes_srcfile[ws_peer_review_minutes_srcfile[0]]
            # 从封面页读取文书番号
            doc_number = ws_cover.cell(row = row_num_document_number, column = col_num_document_number).value
            # pysvn get last changed revision and url
            entry = client.info(peer_reivew_minutes_srcfile)
            # print(MSN[mod_index] + ',' + str(list_peer_review_minutes[i]) + ',' + str(doc_number) + ',' + str(entry.url) + ',' + 'svn:' + str(entry.commit_revision.number))
            print(MSN[mod_index] + ',PeerReviewMinutes,' + entry.url + ',' + str(doc_number) + ',' + 'svn:' + str(entry.commit_revision.number))
            

# *********************************************************************************************************************
# Fetching Safety Plan Attachment Deliveries Products version and last changed revision.
# *********************************************************************************************************************
def FetchSafetyPlanAttachmentDeliveriesProducts():
    print('****************************************************************************************************************')
    print('---------------- Fetching Safety Plan Attachment Deliveries Products last changed revision ---------------------')
    print('****************************************************************************************************************')

    # msn = ['dio','fls','mcu','port','wdg','adc','cortst','eth','flstst','gpt','icu','lin','pwm','ramtst','general','can','spi','fr']
    msn = ['can','spi','fr']

    client = pysvn.Client()

    for index in range(len(msn)):
        # 需要提交的成果物文件路径列表文件，这些文件需要事前根据需要提交的成果物制作，要确保文件的最后一行是空行
        file_path = "U:\\internal\\X1X\\F1x\\common_family\\docs\\FuSa\\Safety_Plan_Ver4.05.00.B_Ver42.05.00.B\\product_deliveries_" + msn[index] + "_path.txt"
        # 打开成果物文件路径列表文件
        wps_deliveries = open(file_path,'r',encoding='utf-8')
        # 遍历成果物文件路径列表文件所有行
        for wps_path_line in wps_deliveries:
            # print(wps_path_line)
            # 检索每行中的回车符并删除
            if "\n" in wps_path_line:
                # 获取成果物文件路径
                wps_path_line = wps_path_line.replace('\n','')
                # print(wps_path_line)
                # 根据成果物文件扩展名对文件进行关键字查询
                if wps_path_line.endswith('.h') or wps_path_line.endswith('.c') or wps_path_line.endswith('.cfgxml') or wps_path_line.endswith('.arxml') or wps_path_line.endswith('.trxml'):
                    # print(wps_path_line)
                    # 如果文件存在
                    if os.path.exists(wps_path_line) == True:
                        # 读取文件信息到wps_file_read
                        with open(wps_path_line,'r',encoding='utf-8') as wps_file_read:
                            # print(wps_file_read)
                            # 创建一个空列表用于保存查询到的记录，最后从列表中打印输出最后一条记录，就是最新的成果物版本号
                            ver_history_list = []
                            # 以行为单位遍历打开的文件内容
                            for line in wps_file_read:
                                # 用正则表达式匹配每一行的内容，确定该行是否包含版本信息(星号，空格，大写字母V，至少一个数字，点，至少一个数字，点，至少一个数字)
                                version_info = re.search(r'\*\sV\d+\.\d+\.\d+',line)
                                # 如果匹配到版本信息
                                if version_info != None:
                                    # 从匹配到的版本信息中把星号和空格去除
                                    version_info = re.search(r'V\d+\.\d+\.\d+',str(version_info.group()))
                                    # 将版本信息追加到之前创建的列表中
                                    ver_history_list.append(version_info.group())
                                else:
                                    pass
                            # 获取文件的SVN revision信息
                            entry = client.info(wps_path_line)
                            # 如果成功匹配到版本号
                            if ver_history_list:
                                # 所有内容检索完成后，打印输出最后一条记录
                                # print(wps_path_line + ',' + str(ver_history_list[len(ver_history_list) - 1]) + '(SVN:' + str(entry.commit_revision.number) + ')' + ',' + str(entry.commit_author))
                                print(wps_path_line + ',' + entry.name + ',' + entry.url + ',' + str(ver_history_list[len(ver_history_list) - 1]) + '(SVN:' + str(entry.commit_revision.number) + ')')
                            # 未成功匹配到版本号(文件版本的记录方式不标准)
                            else:
                                print(wps_path_line + ',' + entry.name + ',' + entry.url + ', Regex Match Failed. Version format is not standard. ' + '(SVN:' + str(entry.commit_revision.number) + ')')
                    # 如果文件不存在
                    else:
                        print(wps_path_line + ',' + entry.name + ',' + entry.url + ', No such file or directory')
                # 空行直接输出空行
                elif wps_path_line == '':
                    print(wps_path_line)
                    # pass
                # 非以上两种情况时需要人工确认版本信息
                else:
                    # 获取文件的SVN revision信息
                    entry = client.info(wps_path_line)
                    print(wps_path_line + ',' + entry.name + ',' + entry.url + ', File of this type need to confirm manually' + '(SVN:' + str(entry.commit_revision.number) + ')')


# *********************************************************************************************************************
# Fetching Safety Plan List of Evidence Coverage Info Snapshot last changed Revision
# snapshot每次都是重新创建，所以每次的文件名都是不同的,需要遍历目录下的所有文件，并抽取最后更新的文件信息
# *********************************************************************************************************************
def FetchCoverageInfoSnapshot():
    print('****************************************************************************************************************')
    print('----------------- Fetching New Added Coverage Info Snapshot Last Changed Revision ------------------------------')
    print('****************************************************************************************************************')

    client = pysvn.Client()
    # the file below defines the path of coverage info snapshot
    path_file_loc = "U:\\internal\\X1X\\F1x\\common_family\\docs\\FuSa\\Safety_Plan_Ver4.05.00.B_Ver42.05.00.B\\coverage_info_path.txt"
    # open file
    file_path = open(path_file_loc,'r',encoding='utf-8')
    # walk through the path listed in the file_path
    for file_path_line in file_path:
        # if file_path_line contains line break
        if "\n" in file_path_line:
            # remove the line break from file_path_line
            file_path_line = file_path_line.replace('\n','')
            # each of the two list below stores the latest two snapshot files path and svn revision
            list_file_path = [0,0]
            list_file_svn_revision = [0,0]
            print('--------------------------------------------------------------------------------')
            try:
                list_file = os.listdir(file_path_line)
                for i in range(len(list_file)):
                    # construct location of the snapshot file
                    file_loc = file_path_line + list_file[i]
                    entry = client.info(file_loc)
                    # compare svn revision and stores the max two revison and corresponding file path
                    if entry.commit_revision.number > list_file_svn_revision[0]:
                        list_file_svn_revision[0] = entry.commit_revision.number
                        list_file_path[0] = file_loc
                    elif entry.commit_revision.number == list_file_svn_revision[0]:
                        list_file_svn_revision[1] = entry.commit_revision.number
                        list_file_path[1] = file_loc
                    else:
                        pass
                print(str(list_file_path[0]) + ',svn: ' + str(list_file_svn_revision[0]))
                print(str(list_file_path[1]) + ',svn: ' + str(list_file_svn_revision[1]))
            except:
                print(file_path_line + ',path incorrect')


# *********************************************************************************************************************
# Safety Analysis Review Checklist last changed revison fetching
# *********************************************************************************************************************
def FetchSafetyAnalysisReviewChecklist():
    MSN = ['DIO','FLS','MCU','PORT','WDG','ADC','CORTST','ETH','FLSTST','GPT','ICU','LIN','PWM','RAMTST']
    msn = ['dio','fls','mcu','port','wdg','adc','cortst','eth','flstst','gpt','icu','lin','pwm','ramtst']

    print('****************************************************************************************************************')
    print('------------------------Safety Analysis Review Checklist Last Modified Revision---------------------------------')
    print('****************************************************************************************************************')

    client = pysvn.Client()

    for index in range(len(MSN)):
        if MSN[index] == 'ETH':
            src_file = "U:\\internal\\X1X\\F1x\\modules\\" + msn[index] + "\\review\\safety_analysis\\Inspection\\F1KM_Ver4.05.00_Ver42.05.00_F1KH_Ver42.05.00_ASILB\\" + MSN[index] + "_Software_Safety_AnalysisChecklist.xlsm"
        else:
            src_file = "U:\\internal\\X1X\\F1x\\modules\\" + msn[index] + "\\review\\safety_analysis\\Inspection\\F1K_F1KM_Ver4.05.00_Ver42.05.00_F1KH_Ver42.05.00_ASILB\\" + MSN[index] + "_Software_Safety_AnalysisChecklist.xlsm"
        entry = client.info(src_file)
        # print(MSN[index] + ',' + src_file + ',svn: ' + str(entry.commit_revision.number))
        print(MSN[index] + ',FMEA Checlist,' + entry.url + ',svn: ' + str(entry.commit_revision.number))

# *********************************************************************************************************************
# Change Management Report version information fetching
# *********************************************************************************************************************
def FetchChangeManagementReport():
    MSN = ['DIO','FLS','MCU','PORT','WDG','ADC','Cortst','ETH','Flstst','GPT','ICU','LIN','PWM','RamTst','General']
    row_version_info = 14
    col_version_info = 3

    print('****************************************************************************************************************')
    print('------------------------------- Change Management Report version information -----------------------------------')
    print('****************************************************************************************************************')

    client = pysvn.Client()

    for index in range(len(MSN)):
        cm_src_file = "U:\\internal\\X1X\\common_platform\\docs\\Impact_analysis\\F1Kx_Ver4.05.00_Ver42.05.00_ASILB\\F1Kx_V4.05.00.B_" + MSN[index] + "_Change_Management.xlsx"
        wb_cm = openpyxl.load_workbook(cm_src_file,data_only = True,read_only = True)
        ws_cm = wb_cm.sheetnames
        ws_cm_cover = wb_cm[ws_cm[0]]
        cm_version = ws_cm_cover.cell(row = row_version_info,column = col_version_info).value

        entry = client.info(cm_src_file)
        print(MSN[index] + ',' + entry.name + ',' + cm_version)


# 当我们在命令行运行模块文件时，Python解释器把一个特殊变量__name__置为__main__,
# 而如果在其他地方导入该模块时，if判断将失败。
if __name__ == '__main__':
    # FetchSafetyAnalysisReviewChecklist()
    # FetchCoverageInfoSnapshot()
    # FetchSafetyPlanAttachmentDeliveriesProducts()
    # FetchPeerReviewMinutes()
    # FetchReqtifyTraceabilityReports()
    # FetchReqtifyUncoveredRequirementsJustificationReport()
    # FetchReqtifyReport()
    # FetchQACReport()
    # FetchTSTRReport()
    # FetchTSTPReport()
    # FetchESTRReport()
    # FetchESTSReport()
    # FetchDFAReport()
    # FetchFMEAReport()
    # FetchTEQReport()
    # FetchUMReport()
    # FetchESDDTSDDReport()
    # FetchUTPUTRReport()
    FetchChangeManagementReport()